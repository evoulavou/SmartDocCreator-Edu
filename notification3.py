"""
notification3.py
----------------
Δημιουργεί ατομικά κοινοποιητήρια (.docx + PDF) για κάθε εκπαιδευτικό
βάσει Excel με πολλαπλά sheets.
"""
from __future__ import annotations

import copy
import os
import re

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from shared import (
    NOTIFICATION_COLUMNS,
    convert_to_pdf,
    has_data,
    load_general_info,
)

ORG_THESIS   = {'s1_Διάθεση οργ.υπεράριθμων', 's2_Διάθεση οργ.τοποθετημένων',
                's4_Απόσπαση εντός'}
DIA_THESIS   = {'s3_Από διάθεση για λειτουργικά', 's6_Τοποθέτηση_διάθεση νεοδιόρισ'}
APOSP_THESIS = {'s5_Πρ_τοποθέτηση αποσπασμένων'}
NO_THESI     = {'s7_Τοποθέτηση_διάθεση αναπληρω'}
BOSS_THESIS  = {'s8_Υπερωρίες σε Δντες'}

INLINE_VERB_SHEETS = (
    's1_Διάθεση οργ.υπεράριθμων',
    's2_Διάθεση οργ.τοποθετημένων',
    's3_Από διάθεση για λειτουργικά',
    's4_Απόσπαση εντός',
    's5_Πρ_τοποθέτηση αποσπασμένων',
    's6_Τοποθέτηση_διάθεση νεοδιόρισ',
    's7_Τοποθέτηση_διάθεση αναπληρω',
    's8_Υπερωρίες σε Δντες',
)

DISPOSITIONS = [
    ('ΣΧΟΛΕΙΟ ΔΙΑΘΕΣΗΣ',    'ΩΡΕΣ ΣΧΟΛΕΙΟΥ ΔΙΑΘΕΣΗΣ'),
    ('ΣΧΟΛΕΙΟ 2ης ΔΙΑΘΕΣΗΣ', 'ΩΡΕΣ 2ου ΣΧΟΛΕΙΟΥ ΔΙΑΘΕΣΗΣ'),
    ('ΣΧΟΛΕΙΟ 3ης ΔΙΑΘΕΣΗΣ', 'ΩΡΕΣ 3ου ΣΧΟΛΕΙΟΥ ΔΙΑΘΕΣΗΣ'),
    ('ΣΧΟΛΕΙΟ 4ης ΔΙΑΘΕΣΗΣ', 'ΩΡΕΣ 4ου ΣΧΟΛΕΙΟΥ ΔΙΑΘΕΣΗΣ'),
]

SCHOOL_COLS_FOR_META = [
    'ΟΡΓΑΝΙΚΗ ΘΕΣΗ', 'ΣΧΟΛΕΙΟ ΑΠΟΣΠΑΣΗΣ', 'ΣΧΟΛΕΙΟ ΑΣΚΗΣΗΣ ΘΗΤΕΙΑΣ ΣΤΕΛΕΧΟΥΣ',
    'ΣΧΟΛΕΙΟ ΤΟΠΟΘΕΤΗΣΗΣ', 'ΣΧΟΛΕΙΟ ΠΡΟΣΩΡΙΝΗΣ ΤΟΠΟΘΕΤΗΣΗΣ',
    'ΣΧΟΛΕΙΟ ΔΙΑΘΕΣΗΣ', 'ΣΧΟΛΕΙΟ 2ης ΔΙΑΘΕΣΗΣ',
    'ΣΧΟΛΕΙΟ 3ης ΔΙΑΘΕΣΗΣ', 'ΣΧΟΛΕΙΟ 4ης ΔΙΑΘΕΣΗΣ',
]

def _build_sheet_texts(hm: str) -> dict:
    """
    Επιστρέφει το dict με τα κείμενα θέσης/διάθεσης για κάθε sheet και filter.
    Το ``hm`` είναι η ημερομηνία απόφασης.

    Κάθε τιμή είναι tuple (τίτλος, σώμα):
      - τίτλος: string ή callable(row) → string
      - σώμα:   callable(row) → string
    """
    def _has_disp(r): return has_data(r.get('ΣΧΟΛΕΙΟ ΔΙΑΘΕΣΗΣ', ''))
    def _h(r, col='ΩΡΕΣ ΘΕΣΗΣ'):
        v = r.get(col)
        return str(v) if has_data(v) else ''

    return {
        's1_Διάθεση οργ.υπεράριθμων': {
            'Εισαγωγή': (
                '«Ανακοίνωση διάθεσης οργανικά υπεράριθμου εκπαιδευτικού για την κάλυψη λειτουργικών αναγκών»',
                lambda r: (
                    f'από το σχολείο της οργανικής σας θέσης ({r["ΟΡΓΑΝΙΚΗ ΘΕΣΗ"]} - {_h(r)} ώρες την εβδομάδα) διατεθήκατε:'
                    if _has_disp(r) else
                    f'από το σχολείο της οργανικής σας θέσης ({r["ΟΡΓΑΝΙΚΗ ΘΕΣΗ"]} - {_h(r)} ώρες την εβδομάδα) από {hm} έως τη λήξη του διδακτικού έτους.'
                ),
            ),
            'Τροποποίηση': (
                '«Ανακοίνωση τροποποίησης διάθεσης οργανικά υπεράριθμου εκπαιδευτικού για την κάλυψη λειτουργικών αναγκών»',
                lambda r: (
                    f'τροποποιήθηκε η διάθεσή σας για την κάλυψη λειτουργικών αναγκών και από το σχολείο της οργανικής σας θέσης ({r["ΟΡΓΑΝΙΚΗ ΘΕΣΗ"]} - {_h(r)} ώρες την εβδομάδα) διατεθήκατε:'
                    if _has_disp(r) else
                    f'τροποποιήθηκε η διάθεσή σας για την κάλυψη λειτουργικών αναγκών και από το σχολείο της οργανικής σας θέσης ({r["ΟΡΓΑΝΙΚΗ ΘΕΣΗ"]} - {_h(r)} ώρες την εβδομάδα) από {hm} έως τη λήξη του διδακτικού έτους.'
                ),
            ),
        },
        's2_Διάθεση οργ.τοποθετημένων': {
            'Εισαγωγή': (
                '«Ανακοίνωση διάθεσης εκπαιδευτικού για την κάλυψη λειτουργικών αναγκών»',
                lambda r: (
                    f'από το σχολείο της οργανικής σας θέσης ({r["ΟΡΓΑΝΙΚΗ ΘΕΣΗ"]} - {_h(r)} ώρες την εβδομάδα) διατεθήκατε:'
                    if _has_disp(r) else
                    f'από το σχολείο της οργανικής σας θέσης ({r["ΟΡΓΑΝΙΚΗ ΘΕΣΗ"]} - {_h(r)} ώρες την εβδομάδα) από {hm} έως τη λήξη του διδακτικού έτους.'
                ),
            ),
            'Τροποποίηση': (
                '«Ανακοίνωση τροποποίησης διάθεσης εκπαιδευτικού από την οργανική του θέση για την κάλυψη λειτουργικών αναγκών»',
                lambda r: (
                    f'τροποποιήθηκε η διάθεσή σας για την κάλυψη λειτουργικών αναγκών και παραμένετε στο σχολείο της οργανικής σας θέσης ({r["ΟΡΓΑΝΙΚΗ ΘΕΣΗ"]} - {_h(r)} ώρες την εβδομάδα) διατεθήκατε:'
                    if _has_disp(r) else
                    f'τροποποιήθηκε η διάθεσή σας για την κάλυψη λειτουργικών αναγκών και παραμένετε στο σχολείο της οργανικής σας θέσης ({r["ΟΡΓΑΝΙΚΗ ΘΕΣΗ"]} - {_h(r)} ώρες την εβδομάδα) από {hm} έως τη λήξη του διδακτικού έτους.'
                ),
            ),
        },
        's3_Από διάθεση για λειτουργικά': {
            'Εισαγωγή': (
                lambda r: (
                    '«Ανακοίνωση τοποθέτησης εκπαιδευτικού, που βρίσκεται στη διάθεση του ΠΥΣΔΕ και διάθεσής του για την κάλυψη λειτουργικών αναγκών»'
                    if _has_disp(r) else
                    '«Ανακοίνωση τοποθέτησης εκπαιδευτικού, που βρίσκεται στη διάθεση του ΠΥΣΔΕ»'
                ),
                lambda r: (
                    f'τοποθετηθήκατε στο {r["ΣΧΟΛΕΙΟ ΠΡΟΣΩΡΙΝΗΣ ΤΟΠΟΘΕΤΗΣΗΣ"]} ({_h(r)} ώρες την εβδομάδα) και διατεθήκατε:'
                    if _has_disp(r) else
                    f'τοποθετηθήκατε στο {r["ΣΧΟΛΕΙΟ ΠΡΟΣΩΡΙΝΗΣ ΤΟΠΟΘΕΤΗΣΗΣ"]} ({_h(r)} ώρες την εβδομάδα) από {hm} έως τη λήξη του διδακτικού έτους.'
                ),
            ),
            'Τροποποίηση': (
                lambda r: (
                    '«Ανακοίνωση τροποποίησης τοποθέτησης εκπαιδευτικού, που βρίσκεται στη διάθεση του ΠΥΣΔΕ και διάθεσής του για την κάλυψη λειτουργικών αναγκών»'
                    if _has_disp(r) else
                    '«Ανακοίνωση τροποποίησης τοποθέτησης εκπαιδευτικού, που βρίσκεται στη διάθεση του ΠΥΣΔΕ»'
                ),
                lambda r: (
                    f'τροποποιήθηκε η τοποθέτηση και η διάθεσή σας, καθώς τοποθετηθήκατε στο {r["ΣΧΟΛΕΙΟ ΠΡΟΣΩΡΙΝΗΣ ΤΟΠΟΘΕΤΗΣΗΣ"]} ({_h(r)} ώρες την εβδομάδα) και διατεθήκατε:'
                    if _has_disp(r) else
                    f'τροποποιήθηκε η τοποθέτησή σας και τοποθετηθήκατε στο {r["ΣΧΟΛΕΙΟ ΠΡΟΣΩΡΙΝΗΣ ΤΟΠΟΘΕΤΗΣΗΣ"]} ({_h(r)} ώρες την εβδομάδα) από {hm} έως τη λήξη του διδακτικού έτους.'
                ),
            ),
        },
        's4_Απόσπαση εντός': {
            'Εισαγωγή': (
                lambda r: (
                    '«Ανακοίνωση απόσπασης εκπαιδευτικού εντός ΠΥΣΔΕ με αίτησή του και διάθεσής του για τη συμπλήρωση του υποχρεωτικού ωραρίου»'
                    if _has_disp(r) else
                    '«Ανακοίνωση απόσπασης εκπαιδευτικού εντός ΠΥΣΔΕ με αίτησή του»'
                ),
                lambda r: (
                    f'αποσπαστήκατε από το σχολείο οργανικής σας θέσης ({r["ΟΡΓΑΝΙΚΗ ΘΕΣΗ"]}) στο {r["ΣΧΟΛΕΙΟ ΑΠΟΣΠΑΣΗΣ"]} ({_h(r)} ώρες την εβδομάδα) και διατεθήκατε:'
                    if _has_disp(r) else
                    f'αποσπαστήκατε από το σχολείο οργανικής σας θέσης ({r["ΟΡΓΑΝΙΚΗ ΘΕΣΗ"]}) στο {r["ΣΧΟΛΕΙΟ ΑΠΟΣΠΑΣΗΣ"]} ({_h(r)} ώρες την εβδομάδα) από {hm} έως τη λήξη του διδακτικού έτους.'
                ),
            ),
            'Τροποποίηση': (
                lambda r: (
                    '«Ανακοίνωση τροποποίησης απόσπασης εκπαιδευτικού εντός ΠΥΣΔΕ με αίτησή του και διάθεσής του για τη συμπλήρωση του υποχρεωτικού ωραρίου»'
                    if _has_disp(r) else
                    '«Ανακοίνωση τροποποίησης απόσπασης εκπαιδευτικού εντός ΠΥΣΔΕ με αίτησή του»'
                ),
                lambda r: (
                    f'τροποποιήθηκε η απόσπαση και διάθεσή σας, καθώς αποσπαστήκατε από το σχολείο οργανικής σας θέσης ({r["ΟΡΓΑΝΙΚΗ ΘΕΣΗ"]}) στο {r["ΣΧΟΛΕΙΟ ΑΠΟΣΠΑΣΗΣ"]} ({_h(r)} ώρες την εβδομάδα) και διατεθήκατε:'
                    if _has_disp(r) else
                    f'τροποποιήθηκε η απόσπασή σας, καθώς αποσπαστήκατε από το σχολείο οργανικής σας θέσης ({r["ΟΡΓΑΝΙΚΗ ΘΕΣΗ"]}) στο {r["ΣΧΟΛΕΙΟ ΑΠΟΣΠΑΣΗΣ"]} ({_h(r)} ώρες την εβδομάδα) από {hm} έως τη λήξη του διδακτικού έτους.'
                ),
            ),
        },
        's5_Πρ_τοποθέτηση αποσπασμένων': {
            'Εισαγωγή': (
                lambda r: (
                    '«Ανακοίνωση τοποθέτησης εκπαιδευτικού, που αποσπάστηκε στο Π.Υ.Σ.Δ.Ε. Χαλκιδικής και διάθεσής του για την κάλυψη λειτουργικών αναγκών»'
                    if _has_disp(r) else
                    '«Ανακοίνωση τοποθέτησης εκπαιδευτικού, που αποσπάστηκε στο Π.Υ.Σ.Δ.Ε. Χαλκιδικής»'
                ),
                lambda r: (
                    f'τοποθετηθήκατε στο {r["ΣΧΟΛΕΙΟ ΠΡΟΣΩΡΙΝΗΣ ΤΟΠΟΘΕΤΗΣΗΣ"]} ({_h(r)} ώρες την εβδομάδα) και:'
                    if _has_disp(r) else
                    f'τοποθετηθήκατε στο {r["ΣΧΟΛΕΙΟ ΠΡΟΣΩΡΙΝΗΣ ΤΟΠΟΘΕΤΗΣΗΣ"]} ({_h(r)} ώρες την εβδομάδα) από {hm} έως τη λήξη του διδακτικού έτους.'
                ),
            ),
            'Τροποποίηση': (
                lambda r: (
                    '«Ανακοίνωση τροποποίησης τοποθέτησης εκπαιδευτικού, που αποσπάστηκε στο Π.Υ.Σ.Δ.Ε. Χαλκιδικής και διάθεσής του για την κάλυψη λειτουργικών αναγκών»'
                    if _has_disp(r) else
                    '«Ανακοίνωση τροποποίησης τοποθέτησης εκπαιδευτικού, που αποσπάστηκε στο Π.Υ.Σ.Δ.Ε. Χαλκιδικής»'
                ),
                lambda r: (
                    f'τροποποιήθηκε η τοποθέτηση και η διάθεσή σας, καθώς τοποθετηθήκατε στο {r["ΣΧΟΛΕΙΟ ΠΡΟΣΩΡΙΝΗΣ ΤΟΠΟΘΕΤΗΣΗΣ"]} ({_h(r)} ώρες την εβδομάδα) και διατεθήκατε:'
                    if _has_disp(r) else
                    f'τροποποιήθηκε η τοποθέτησή σας και τοποθετηθήκατε στο {r["ΣΧΟΛΕΙΟ ΠΡΟΣΩΡΙΝΗΣ ΤΟΠΟΘΕΤΗΣΗΣ"]} ({_h(r)} ώρες την εβδομάδα) από {hm} έως τη λήξη του διδακτικού έτους.'
                ),
            ),
        },
        's6_Τοποθέτηση_διάθεση νεοδιόρισ': {
            'Εισαγωγή': (
                lambda r: (
                    '«Ανακοίνωση τοποθέτησης νεοδιόριστου εκπαιδευτικού και διάθεσής του για την κάλυψη λειτουργικών αναγκών»'
                    if _has_disp(r) else
                    '«Ανακοίνωση τοποθέτησης νεοδιόριστου εκπαιδευτικού»'
                ),
                lambda r: (
                    f'τοποθετηθήκατε στο {r["ΣΧΟΛΕΙΟ ΠΡΟΣΩΡΙΝΗΣ ΤΟΠΟΘΕΤΗΣΗΣ"]} και διατεθήκατε: '
                    if _has_disp(r) else
                    f'τοποθετηθήκατε στο {r["ΣΧΟΛΕΙΟ ΠΡΟΣΩΡΙΝΗΣ ΤΟΠΟΘΕΤΗΣΗΣ"]} από {hm} έως τη λήξη του διδακτικού έτους.'
                ),
            ),
            'Τροποποίηση': (
                lambda r: (
                    '«Ανακοίνωση τροποποίησης τοποθέτησης νεοδιόριστου εκπαιδευτικού και διάθεσής του για την κάλυψη λειτουργικών αναγκών»'
                    if _has_disp(r) else
                    '«Ανακοίνωση τροποποίησης τοποθέτησης νεοδιόριστου εκπαιδευτικού»'
                ),
                lambda r: (
                    f'τροποποιήθηκε η τοποθέτηση και η διάθεσή σας, καθώς τοποθετηθήκατε στο {r["ΣΧΟΛΕΙΟ ΠΡΟΣΩΡΙΝΗΣ ΤΟΠΟΘΕΤΗΣΗΣ"]} και διατεθήκατε:'
                    if _has_disp(r) else
                    f'τροποποιήθηκε η τοποθέτησή σας και τοποθετηθήκατε στο {r["ΣΧΟΛΕΙΟ ΠΡΟΣΩΡΙΝΗΣ ΤΟΠΟΘΕΤΗΣΗΣ"]} από  {hm} έως τη λήξη του διδακτικού έτους.'
                ),
            ),
        },
        's7_Τοποθέτηση_διάθεση αναπληρω': {
            'Εισαγωγή': (
                lambda r: (
                    '«Ανακοίνωση τοποθέτησης αναπληρωτή εκπαιδευτικού γενικής εκπαίδευσης και διάθεσής του για την κάλυψη λειτουργικών αναγκών»'
                    if _has_disp(r) else
                    '«Ανακοίνωση τοποθέτησης αναπληρωτή εκπαιδευτικού γενικής εκπαίδευσης»'
                ),
                lambda r: (
                    f'τοποθετηθήκατε στο {r["ΣΧΟΛΕΙΟ ΤΟΠΟΘΕΤΗΣΗΣ"]} ({r["ΩΡΕΣ"]} ώρες την εβδομάδα) και διατεθήκατε:'
                    if _has_disp(r) else
                    f'τοποθετηθήκατε στο {r["ΣΧΟΛΕΙΟ ΤΟΠΟΘΕΤΗΣΗΣ"]} από  {hm} έως τη λήξη του διδακτικού έτους.'
                ),
            ),
            'Τροποποίηση': (
                lambda r: (
                    '«Ανακοίνωση τροποποίησης τοποθέτησης αναπληρωτή εκπαιδευτικού γενικής εκπαίδευσης και διάθεσής του για την κάλυψη λειτουργικών αναγκών»'
                    if _has_disp(r) else
                    '«Ανακοίνωση τροποποίησης τοποθέτησης αναπληρωτή εκπαιδευτικού γενικής εκπαίδευσης»'
                ),
                lambda r: (
                    f'τροποποιήθηκε η τοποθέτηση και η διάθεσή σας, καθώς τοποθετηθήκατε στο {r["ΣΧΟΛΕΙΟ ΤΟΠΟΘΕΤΗΣΗΣ"]} ({r["ΩΡΕΣ"]} ώρες την εβδομάδα) και διατεθήκατε:'
                    if _has_disp(r) else
                    f'τροποποιήθηκε η τοποθέτησή σας, καθώς τοποθετηθήκατε στο {r["ΣΧΟΛΕΙΟ ΤΟΠΟΘΕΤΗΣΗΣ"]} από  {hm} έως τη λήξη του διδακτικού έτους.'
                ),
            ),
        },
        's8_Υπερωρίες σε Δντες': {
            'Εισαγωγή': (
                lambda r: '«Ανακοίνωση ανάθεσης υπερωριακής εργασίας, για την κάλυψη λειτουργικών αναγκών, σε διευθυντή σχολικής μονάδας»',
                lambda r: f'σας ανατίθεται υπερωριακή εργασία {r["ΩΡΕΣ ΥΠΕΡΩΡΙΑΣ"]} ωρών για την κάλυψη λειτουργικών αναγκών στο {r["ΣΧΟΛΕΙΟ ΑΣΚΗΣΗΣ ΘΗΤΕΙΑΣ ΣΤΕΛΕΧΟΥΣ"]}, όπου υπηρετείτε με θητεία ως διευθυντής, από {hm} έως τη λήξη του διδακτικού έτους.',
            ),
            'Τροποποίηση': (
                lambda r: '«Ανακοίνωση τροποποίησης ανάθεσης υπερωριακής εργασίας σε διευθυντή σχολικής μονάδας για την κάλυψη λειτουργικών αναγκών»',
                lambda r: f'τροποποιείται από {hm} η ανάθεση υπερωριακής εργασίας και σας ανατίθεται υπερωριακή απασχόληση {r["ΩΡΕΣ ΥΠΕΡΩΡΙΑΣ"]} ωρών για την κάλυψη λειτουργικών αναγκών στο {r["ΣΧΟΛΕΙΟ ΑΣΚΗΣΗΣ ΘΗΤΕΙΑΣ ΣΤΕΛΕΧΟΥΣ"]}, όπου υπηρετείτε με θητεία ως διευθυντής, έως τη λήξη του διδακτικού έτους.',
            ),
        },
    }

def _parse_hours(hours_val) -> tuple[int | None, int | None]:
    """
    Αναλύει τιμές όπως '3+5Υ'.
    Επιστρέφει (συνολικές, υπερωριακές) ή (None, None).
    """
    match = re.match(r'^(\d+)\+(\d+)[Υυ]$', str(hours_val).strip())
    if match:
        regular, overtime = int(match.group(1)), int(match.group(2))
        return regular + overtime, overtime
    return None, None

def _apply_font(run, bold: bool = False) -> None:
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = bold


def _replace_general_info(doc: Document, replacements: dict[str, str]) -> None:
    """Αντικαθιστά [Ημερομηνία απόφασης], [Αριθμός...] κλπ."""
    for term, value in replacements.items():
        placeholder = f'[{term}]'
        for p in doc.paragraphs:
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, value)
                if p.runs:
                    _apply_font(p.runs[0], bold=True)


def _get_am(row: pd.Series) -> str:
    for k in ('Α.Μ.', 'Α.Μ', 'AM'):
        if k in row and has_data(row[k]):
            return str(row[k])
    return '-'


def _replace_personal_info(doc: Document, row: pd.Series, key: str, am_val: str) -> None:
    """Αντικαθιστά [Επώνυμο], [Όνομα], [Α.Μ.], [Υπ. Ωράριο], [Τωρινή Θέση]."""
    for el in doc.element.iter():
        if el.tag != qn('w:t'):
            continue
        el.text = el.text.replace('[Επώνυμο]', str(row['ΕΠΩΝΥΜΟ']))
        el.text = el.text.replace('[Όνομα]', str(row['ΟΝΟΜΑ']))
        el.text = el.text.replace('[Ειδικότητα]', str(row['ΕΙΔΙΚΟΤΗΤΑ']))
        el.text = el.text.replace('[Α.Μ.]', am_val)

        wr = '-' if key == 's6_Τοποθέτηση_διάθεση νεοδιόρισ' else str(row.get('ΥΠΟΧΡΕΩΤΙΚΟ ΩΡΑΡΙΟ', '-'))
        el.text = el.text.replace('[Υπ. Ωράριο]', wr)

        if '[Τωρινή Θέση]' in el.text:
            if key in ORG_THESIS:
                el.text = el.text.replace('[Τωρινή Θέση]', f'Οργανική Θέση: {row["ΟΡΓΑΝΙΚΗ ΘΕΣΗ"]}')
            elif key in DIA_THESIS:
                el.text = el.text.replace('[Τωρινή Θέση]', 'Στη Διάθεση του ΠΥΣΔΕ Χαλκιδικής')
            elif key in APOSP_THESIS:
                el.text = el.text.replace('[Τωρινή Θέση]', f'Με Απόσπαση απο: {row["ΔΔΕ/ΔΠΕ ΟΡΓΑΝΙΚΗΣ ΘΕΣΗΣ"]}')
            elif key in NO_THESI:
                el.text = el.text.replace('[Τωρινή Θέση]', 'Τωρινή θέση: -')
            elif key in BOSS_THESIS:
                el.text = el.text.replace('[Τωρινή Θέση]', f'Σχολείο θητείας: {row["ΣΧΟΛΕΙΟ ΑΣΚΗΣΗΣ ΘΗΤΕΙΑΣ ΣΤΕΛΕΧΟΥΣ"]}')


def _replace_dde(doc: Document, row: pd.Series, key: str) -> None:
    """Αντικαθιστά [ΔΔΕ] με το σχολείο απόσπασης ή '-'."""
    value = str(row['ΔΔΕ/ΔΠΕ ΟΡΓΑΝΙΚΗΣ ΘΕΣΗΣ']) if key in APOSP_THESIS else '-'
    for p in doc.paragraphs:
        if '[ΔΔΕ]' in p.text:
            p.text = p.text.replace('[ΔΔΕ]', value)
            if p.runs:
                _apply_font(p.runs[0], bold=True)

def _build_disposition_texts(row: pd.Series, key: str, is_inline: bool) -> list[str]:
    """
    Δημιουργεί τη λίστα κειμένων για τα bullets διάθεσης.
    Για inline sheets (s1/s2 κλπ) το ρήμα αγνοείται.
    """
    texts = []
    for school_col, hours_col in DISPOSITIONS:
        if school_col not in row or not has_data(row[school_col]):
            continue
        school_val = row[school_col]
        hours_val  = row.get(hours_col, '')
        total_h, overtime_h = _parse_hours(hours_val)

        if is_inline:
            if overtime_h is not None:
                texts.append(f'στο {school_val} για {total_h} ώρες την εβδομάδα (εκ των οποίων {overtime_h} υπερωριακά)')
            else:
                texts.append(f'στο {school_val} για {hours_val} ώρες την εβδομάδα')
        else:
            if overtime_h is not None:
                texts.append(f'διατεθήκατε στο {school_val} για {total_h} ώρες την εβδομάδα (εκ των οποίων {overtime_h} υπερωριακά)')
            else:
                texts.append(f'διατεθήκατε στο {school_val} για {hours_val} ώρες την εβδομάδα')
    return texts


def _format_bullets(items: list[str], prepend_first: str | None) -> list[str]:
    """
    Προσθέτει ',' / 'και' στα bullets και προαιρετικά προτάσσει ένα extra πρώτο item.
    """
    all_items = ([prepend_first] + items) if prepend_first else list(items)
    if not all_items:
        return []
    result = []
    for i, text in enumerate(all_items):
        is_second_to_last = (len(all_items) > 1) and (i == len(all_items) - 2)
        result.append(text + (' και' if is_second_to_last else ','))
    return result


def _insert_bullets(doc: Document, bullet_paragraph, bullet_idx: int,
                    formatted_bullets: list[str], final_date_text: str) -> None:
    """Εισάγει τα bullets στο document και προσθέτει την ημερομηνία στο τέλος."""

    bullet_paragraph.text = bullet_paragraph.text.replace('[Θέση1_Bullet]', formatted_bullets[0])
    if bullet_paragraph.runs:
        _apply_font(bullet_paragraph.runs[0])

    current_idx = bullet_idx + 1

    for f_text in formatted_bullets[1:]:
        if current_idx < len(doc.paragraphs):
            new_p = doc.paragraphs[current_idx].insert_paragraph_before('')
        else:
            new_p = doc.add_paragraph('')

        new_p.style = bullet_paragraph.style
        if bullet_paragraph._p.pPr is not None:
            new_p_pr = new_p._p.get_or_add_pPr()
            new_p_pr.clear()
            for el in bullet_paragraph._p.pPr:
                new_p_pr.append(copy.deepcopy(el))

        r = new_p.add_run(f_text)
        _apply_font(r)
        current_idx += 1

    # Ημερομηνία
    if current_idx < len(doc.paragraphs):
        new_p_final = doc.paragraphs[current_idx].insert_paragraph_before('')
    else:
        new_p_final = doc.add_paragraph('')

    try:
        new_p_final.style = doc.styles['Normal']
    except KeyError:
        pass

    _apply_font(new_p_final.add_run(final_date_text))


def _replace_schools_paragraph(doc: Document, found_schools: list[str]) -> None:
    """Αντικαθιστά [Σχολεία] με πολλαπλά paragraphs (ένα ανά σχολείο)."""
    for p in doc.paragraphs:
        if '[Σχολεία]' not in p.text:
            continue
        parent   = p._element.getparent()
        position = list(parent).index(p._element)

        if found_schools:
            for i, school in enumerate(found_schools):
                clone = copy.deepcopy(p._element)
                for t in clone.iter(qn('w:t')):
                    t.text = t.text.replace('[Σχολεία]', school)
                parent.insert(position + i, clone)
            parent.remove(p._element)
        else:
            p.text = '-'
        break

def notifications3(filename: str) -> None:
    df = pd.read_excel(filename, sheet_name=None, header=None)
    general_info = load_general_info(df)
    hm = general_info.get('Ημερομηνία απόφασης', '....')
    sheet_texts = _build_sheet_texts(hm)

    valid_filters = {'Εισαγωγή', 'Τροποποίηση', 'Ανάκληση'}

    for key, col_names in NOTIFICATION_COLUMNS.items():
        if key not in df:
            continue
        print(key)

        data = df[key]
        new_df = data[data.iloc[:, 0].isin(valid_filters)].copy()
        new_df.columns = col_names
        is_inline = key in INLINE_VERB_SHEETS

        for _, row in new_df.iterrows():
            filter_val = str(row['Filter']).strip()

            if key not in sheet_texts or filter_val not in sheet_texts[key]:
                print(f'  Παράλειψη: δεν υπάρχει "{filter_val}" για "{key}"')
                continue

            doc    = Document('TEMPLATE_ΚΟΙΝΟΠΟΙΗΤΗΡΙΑ.docx')
            am_val = _get_am(row)

            _replace_general_info(doc, general_info)
            _replace_personal_info(doc, row, key, am_val)
            _replace_dde(doc, row, key)

            title_fn, body_fn = sheet_texts[key][filter_val]
            txt_thesi  = title_fn(row) if callable(title_fn) else title_fn
            txt_thesi1 = body_fn(row)
            has_disp   = has_data(row.get('ΣΧΟΛΕΙΟ ΔΙΑΘΕΣΗΣ', ''))

            # Αντικατάσταση [Θέση], [Θέση1], [κατηγορία]
            for p in doc.paragraphs:
                if '[Θέση]' in p.text:
                    p.text = p.text.replace('[Θέση]', txt_thesi)
                    if p.runs:
                        _apply_font(p.runs[0], bold=True)

                if '[Θέση1]' in p.text:
                    replacement = txt_thesi1
                    p.text = p.text.replace('[Θέση1]', replacement)
                    if p.runs:
                        _apply_font(p.runs[0])

                if '[κατηγορία]' in p.text:
                    p.text = p.text.replace('[κατηγορία]', 'που στηρίχθηκε')
                    if p.runs:
                        _apply_font(p.runs[0])

            # Εύρεση bullet placeholder
            bullet_paragraph = None
            bullet_idx = -1
            for idx, p in enumerate(doc.paragraphs):
                if '[Θέση1_Bullet]' in p.text:
                    bullet_paragraph, bullet_idx = p, idx
                    break

            if has_disp and bullet_paragraph is not None:
                disp_texts      = _build_disposition_texts(row, key, is_inline)
                prepend         = None if is_inline else txt_thesi1
                formatted       = _format_bullets(disp_texts, prepend)
                final_date_text = f'από {hm} έως τη λήξη του διδακτικού έτους.'
                _insert_bullets(doc, bullet_paragraph, bullet_idx, formatted, final_date_text)

            elif not has_disp and bullet_paragraph is not None:
                elem = bullet_paragraph._element
                elem.getparent().remove(elem)

            # Σχολεία metadata
            found_schools = [str(row[col]) for col in SCHOOL_COLS_FOR_META
                             if col in row and has_data(row[col])]
            _replace_schools_paragraph(doc, found_schools)

            if found_schools:
                schools_str = ';'.join(found_schools)
                doc.core_properties.keywords = ', '.join(found_schools)
                doc.core_properties.subject  = 'Notification regarding: ' + schools_str

            os.makedirs('notifications', exist_ok=True)
            doc.save(
                f'notifications/notification_{key}_for_{row["ΕΠΩΝΥΜΟ"]}_{am_val}_{filter_val}.docx'
            )

    convert_to_pdf('notifications/')