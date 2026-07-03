"""
placement.py
------------
Δημιουργεί το έγγραφο τοποθετήσεων (TEMPLATE_ΤΟΠΟΘΕΤΗΣΕΙΣ.docx)
με πίνακες ανά sheet και status.
"""
from __future__ import annotations

from collections import defaultdict

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from shared import PLACEMENT_COLUMNS, has_data, load_general_info

SHEET_TEXTS: dict[str, dict[str, str]] = {
    's1_Διάθεση οργ.υπεράριθμων': {
        'Εισαγωγή':    'Οι παρακάτω οργανικά υπεράριθμοι εκπαιδευτικοί διατίθενται, για τη συμπλήρωση του υποχρεωτικού ωραρίου, ως εξής:',
        'Τροποποίηση': 'Τροποποιείται η διάθεση των παρακάτω οργανικά υπεράριθμων εκπαιδευτικών, για τη συμπλήρωση του υποχρεωτικού ωραρίου, ως εξής:',
    },
    's2_Διάθεση οργ.τοποθετημένων': {
        'Εισαγωγή':    'Οι παρακάτω εκπαιδευτικοί διατίθενται από την οργανική τους θέση, για τη συμπλήρωση τους υποχρεωτικού ωραρίου, ως εξής:',
        'Τροποποίηση': 'Τροποποιείται η διάθεση των παρακάτω εκπαιδευτικών από την οργανική τους θέση, για τη συμπλήρωση τους υποχρεωτικού ωραρίου, ως εξής:',
    },
    's3_Από διάθεση για λειτουργικά': {
        'Εισαγωγή':    'Οι παρακάτω εκπαιδευτικοί, που ανήκουν στη διάθεση του Π.Υ.Σ.Δ.Ε. Χαλκιδικής, τοποθετούνται και διατίθενται κατά περίπτωση, για την κάλυψη λειτουργικών αναγκών στις παρακάτω Σχολικές Μονάδες, ως εξής:',
        'Τροποποίηση': 'Τροποποιείται η τοποθέτηση και η διάθεση κατά περίπτωση, για την κάλυψη λειτουργικών αναγκών, των παρακάτω εκπαιδευτικών, που ανήκουν στη διάθεση του Π.Υ.Σ.Δ.Ε. Χαλκιδικής, ως εξής:',
    },
    's4_Απόσπαση εντός': {
        'Εισαγωγή':    'Οι παρακάτω εκπαιδευτικοί αποσπώνται, με αίτησή τους, από την οργανική τους θέση και διατίθενται κατά περίπτωση, για την κάλυψη λειτουργικών αναγκών, στις παρακάτω Σχολικές Μονάδες, ως εξής:',
        'Τροποποίηση': 'Τροποποιείται η απόσπαση, με αίτησή τους, από την οργανική τους θέση των παρακάτω εκπαιδευτικών και η διάθεσή τους κατά περίπτωση, για την κάλυψη λειτουργικών αναγκών, στις παρακάτω Σχολικές Μονάδες, ως εξής:',
    },
    's5_Πρ_τοποθέτηση αποσπασμένων': {
        'Εισαγωγή':    'Οι παρακάτω εκπαιδευτικοί που αποσπάστηκαν στο Π.Υ.Σ.Δ.Ε. Χαλκιδικής τοποθετούνται και διατίθενται κατά περίπτωση, για την κάλυψη λειτουργικών αναγκών στις παρακάτω Σχολικές Μονάδες, ως εξής:',
        'Τροποποίηση': 'Τροποποιείται η τοποθέτηση των παρακάτω εκπαιδευτικών που αποσπάστηκαν στο Π.Υ.Σ.Δ.Ε. Χαλκιδικής και η διάθεσή σους κατά περίπτωση, για την κάλυψη λειτουργικών αναγκών στις παρακάτω Σχολικές Μονάδες, ως εξής:',
    },
    's6_Τοποθέτηση_διάθεση νεοδιόρισ': {
        'Εισαγωγή':    'Οι παρακάτω νεοδιόριστοι εκπαιδευτικοί τοποθετούνται και διατίθενται κατά περίπτωση, για την κάλυψη λειτουργικών αναγκών στις παρακάτω Σχολικές Μονάδες, ως εξής:',
        'Τροποποίηση': 'Τροποποιείται η τοποθέτηση και η διάθεση κατά περίπτωση, για την κάλυψη λειτουργικών αναγκών, των παρακάτω νεοδιόριστων εκπαιδευτικών, ως εξής:',
    },
    's7_Τοποθέτηση_διάθεση αναπληρω': {
        'Εισαγωγή':    'Οι παρακάτω αναπληρωτές εκπαιδευτικοί γενικής εκπαίδευσης τοποθετούνται και διατίθενται κατά περίπτωση, για την κάλυψη λειτουργικών αναγκών στις παρακάτω Σχολικές Μονάδες, ως εξής:',
        'Τροποποίηση': 'Τροποποιείται η τοποθέτηση και η διάθεση κατά περίπτωση, για την κάλυψη λειτουργικών αναγκών, των παρακάτω αναπληρωτών εκπαιδευτικών γενικής εκπαίδευσης, ως εξής:',
    },
    's8_Υπερωρίες σε Δντες': {
        'Εισαγωγή':    'Στους παρακάτω Διευθυντές σχολικών μονάδων ανατίθεται υπερωριακή εργασία, για την κάλυψη λειτουργικών αναγκών, ως εξής:',
        'Τροποποίηση': 'Τροποποιείται η ανάθεση υπερωριακής εργασίας, για την κάλυψη λειτουργικών αναγκών, στους παρακάτω Διευθυντές σχολικών μονάδων, ως εξής:',
    },
}

STATUSES = ('Εισαγωγή', 'Τροποποίηση', 'Ανάκληση')

SCHOOL_LABELS = {
    'ΟΡΓΑΝΙΚΗ ΘΕΣΗ', 'ΣΧΟΛΕΙΟ ΑΠΟΣΠΑΣΗΣ', 'ΣΧΟΛΕΙΟ ΔΙΑΘΕΣΗΣ',
    'ΣΧΟΛΕΙΟ 2ης ΔΙΑΘΕΣΗΣ', 'ΣΧΟΛΕΙΟ 3ης ΔΙΑΘΕΣΗΣ',
    'ΣΧΟΛΕΙΟ 4ης ΔΙΑΘΕΣΗΣ', 'ΣΧΟΛΕΙΟ ΠΡΟΣΩΡΙΝΗΣ ΤΟΠΟΘΕΤΗΣΗΣ',
}

def _numbered_headers(header_series) -> list[str]:
    """
    Μετονομάζει διπλότυπες επικεφαλίδες με αύξοντα αριθμό.
    π.χ. ΩΡΕΣ, ΩΡΕΣ, ΩΡΕΣ → ΩΡΕΣ1, ΩΡΕΣ2, ΩΡΕΣ3
    """
    count: dict[str, int] = defaultdict(int)
    for h in header_series:
        count[str(h).strip()] += 1

    seen: dict[str, int] = defaultdict(int)
    result = []
    for h in header_series:
        name = str(h).strip()
        if count[name] > 1:
            seen[name] += 1
            result.append(f'{name}{seen[name]}')
        else:
            result.append(name)
    return result


def _is_hours_col(name: str) -> bool:
    return name.startswith('ΩΡΕΣ') and name[4:].isdigit()


def _resolve_columns(wanted_columns: list[str], header_pos: dict[str, int],
                     new_df: pd.DataFrame) -> tuple[list, list, list]:
    """
    Αντιστοιχεί τα wanted_columns στις πραγματικές θέσεις του DataFrame,
    αφαιρώντας ζεύγη ΣΧΟΛΕΙΟ+ΩΡΕΣ που είναι εντελώς κενά.

    Επιστρέφει (final_positions, final_labels, final_raw).
    """
    pd.set_option('future.no_silent_downcasting', True)

    # Πρώτο πέρασμα: αντιστοίχιση
    selected_positions: list = []
    selected_labels:    list = []
    selected_raw:       list = []

    for col_name in wanted_columns:
        if col_name == 'Α/Α':
            selected_positions.append(None)
            selected_labels.append('Α/Α')
            selected_raw.append('Α/Α')
            continue
        if col_name not in header_pos:
            print(f"Warning: column '{col_name}' not found in sheet, skipping.")
            continue
        selected_positions.append(header_pos[col_name])
        selected_raw.append(col_name)
        display = 'ΩΡΕΣ' if _is_hours_col(col_name) else col_name
        selected_labels.append(display)

    final_positions: list = []
    final_labels:    list = []
    i = 0
    while i < len(selected_positions):
        pos   = selected_positions[i]
        label = selected_labels[i]
        raw   = selected_raw[i]

        if pos is None:
            final_positions.append(None)
            final_labels.append(label)
            i += 1
            continue

        col_data = new_df.iloc[:, pos].fillna('-').infer_objects(copy=False)
        is_empty = (col_data == '-').all()

        if label in SCHOOL_LABELS and is_empty:
            i += 1
            if i < len(selected_raw) and _is_hours_col(selected_raw[i]):
                i += 1
            continue

        final_positions.append(pos)
        final_labels.append(label)
        i += 1

    return final_positions, final_labels, []


def _create_filtered_table(doc: Document, wanted_columns: list[str],
                           sheet_data: pd.DataFrame, status_filter: str):
    """
    Δημιουργεί και επιστρέφει έναν Word πίνακα για τις γραμμές με status == status_filter.
    Επιστρέφει None αν δεν υπάρχουν δεδομένα.
    """
    header_row = sheet_data.iloc[0]
    data_rows  = sheet_data.iloc[1:]

    new_df = data_rows[data_rows.iloc[:, 0] == status_filter].copy()
    new_df = new_df.iloc[:, 1:]
    excel_headers_raw = header_row.iloc[1:]

    if new_df.empty:
        return None

    excel_headers_numbered = _numbered_headers(excel_headers_raw)
    header_pos = {name: pos for pos, name in enumerate(excel_headers_numbered)}

    final_positions, final_labels, _ = _resolve_columns(wanted_columns, header_pos, new_df)

    if not final_labels:
        return None

    num_data_rows = len(new_df)
    table = doc.add_table(rows=num_data_rows + 1, cols=len(final_labels))
    table.style = 'Table Grid'

    # Επικεφαλίδες
    for i, col_name in enumerate(final_labels):
        table.rows[0].cells[i].text = str(col_name)

    # Δεδομένα
    for i in range(num_data_rows):
        for col_idx, pos in enumerate(final_positions):
            if pos is None:
                text_val = str(i + 1)
            else:
                val = new_df.iloc[i, pos]
                text_val = '-' if pd.isna(val) else str(val)
            table.rows[i + 1].cells[col_idx].text = text_val

    return table

def placements(filename: str) -> None:
    df           = pd.read_excel(filename, sheet_name=None, header=None)
    general_info = load_general_info(df)

    doc = Document('TEMPLATE_ΤΟΠΟΘΕΤΗΣΕΙΣ.docx')

    for term, value in general_info.items():
        placeholder = f'[{term}]'
        for p in doc.paragraphs:
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, value)

    anchor_paragraph = next(
        (p for p in doc.paragraphs if '[sheet]' in p.text), None
    )
    if anchor_paragraph is None:
        doc.save('Τοποθέτηση.docx')
        return

    current_element   = anchor_paragraph._element
    is_first_insertion = True
    num = 0

    for key in df:
        if key == 'Στοιχεία απόφασης':
            continue
        if key not in SHEET_TEXTS or key not in PLACEMENT_COLUMNS:
            continue

        texts_for_key = SHEET_TEXTS[key]

        for status in STATUSES:
            data_rows  = df[key].iloc[1:]
            df_subset  = data_rows[data_rows.iloc[:, 0] == status]

            if df_subset.empty:
                continue

            status_text = texts_for_key.get(status, '')
            if not status_text:
                continue

            num += 1

            if is_first_insertion:
                anchor_paragraph.text = f'{num}. {status_text}'
                is_first_insertion = False
            else:
                new_p = doc.add_paragraph(f'{num}. {status_text}')
                current_element.addnext(new_p._element)
                current_element = new_p._element
                new_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            table = _create_filtered_table(doc, PLACEMENT_COLUMNS[key], df[key], status)
            if table:
                current_element.addnext(table._element)
                current_element = table._element

            spacer = doc.add_paragraph()
            current_element.addnext(spacer._element)
            current_element = spacer._element

    if is_first_insertion and '[sheet]' in anchor_paragraph.text:
        anchor_paragraph.text = ''

    doc.save('Τοποθέτηση.docx')